from docx import Document
from docx.shared import Inches, Pt, RGBColor


from sqlalchemy.orm import Session
from sqlalchemy import and_, func
from io import BytesIO
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from database import Session, LMAPrediction, MagneticData, PictureDefect,Event5, RopeDetails

def get_event_with_picture_data(session: Session, date_str: str):
    # 查询 Event5 表中与 date_str 匹配的数据
    event5_data = session.query(Event5).filter(Event5.date == date_str).all()

    # 初始化结果列表
    results = []

    for event in event5_data:
        print(f"Original position: {event.position}")

        # 使用 ROUND() 函数对 position 四舍五入，确保精度一致
        rounded_position = round(event.position, 2)  # 四舍五入到两位小数

        # 查询 PictureDefect 中与（date, position）匹配的数据，四舍五入 position
        picture_defect_data = session.query(PictureDefect).filter(
            and_(
                PictureDefect.date == event.date,
                func.round(PictureDefect.position, 2) == rounded_position  # 对 position 四舍五入
            )
        ).first()  # 假设每个 (date, position) 在 PictureDefect 中唯一

        if picture_defect_data:
            print("Found matching PictureDefect data!")
            result = {
                "date": event.date,
                "position": event.position,
                "LMA_value": event.LMA_value,
                "diameter": event.diameter,
                "defect_degree": event.defect_degree,
                "defect_type": event.defect_type,
                "picture_data": picture_defect_data.picture_data,
                "issue": picture_defect_data.issue,
            }
        else:
            result = {
                "date": event.date,
                "position": event.position,
                "LMA_value": event.LMA_value,
                "diameter": event.diameter,
                "defect_degree": event.defect_degree,
                "defect_type": event.defect_type,
                "picture_data": None,
                "issue": None,
            }

        results.append(result)

    return results


# 根据LMA值分类并标记勾选
def classify_and_check_lma(LMA_value):
    # 根据LMA值确定勾选的列
    categories = ['轻微', '轻度', '中度', '较重', '严重', '超限']
    checks = [False] * 6  # 默认所有勾选项都是False

    if LMA_value >= 1.0:  # 超限损伤
        checks[5] = True
    elif LMA_value >= 0.8:  # 严重损伤
        checks[4] = True
    elif LMA_value >= 0.6:  # 较重损伤
        checks[3] = True
    elif LMA_value >= 0.4:  # 中度损伤
        checks[2] = True
    elif LMA_value >= 0.2:  # 轻度损伤
        checks[1] = True
    else:  # 轻微损伤
        checks[0] = True

    return checks


# 插入带图片的表格到Word模板
def insert_table_with_images(doc, data):
    para = doc.add_paragraph()
    run = para.add_run("重要损伤")
    run.bold = True  # 加粗
    run.font.size = Pt(20)  # 设置大字体
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
    table = doc.add_table(rows=1, cols=6)  # 创建表格

    for record in data:
        # 创建新表格，有4行6列
        table = doc.add_table(rows=4, cols=6)
        table.style = 'Table Grid'

        # 填写表格的第一列
        table.cell(0, 0).text = '绳号'
        table.cell(1, 0).text = '位置'
        table.cell(2, 0).text = '量值'
        table.cell(3, 0).text = '等级'

        # 填写默认的和从记录中获取的数据
        table.cell(0, 1).text = '1'  # 绳号默认值为1
        table.cell(1, 1).text = str(record['position'])
        table.cell(2, 1).text = str(record['LMA_value'])
        table.cell(3, 1).text = str(record['defect_degree'])

        # 填写第五和第六列的标题
        table.cell(3, 2).text = '公称直径/mm'
        table.cell(3, 4).text = '实测直径/mm'

        # 填写直径数据
        table.cell(3, 3).text = '50.5'  # 公称直径默认值为50.5
        table.cell(3, 5).text = str(record['diameter'])

        # 合并单元格用于插入图片
        a = table.cell(0, 2)
        b = table.cell(2, 5)
        a.merge(b)

        # 插入图片到合并后的单元格中
        if record['picture_data']:
            print(f"Picture data length: {len(record['picture_data'])} bytes")

            # 将二进制图片数据转换为内存中的文件对象
            img_stream = BytesIO(record['picture_data'])

            # 插入图片到合并后的单元格中
            merged_cell = table.cell(0, 2)
            paragraph = merged_cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(img_stream, width=Inches(2.5))

        # 在插入一个表格之后，添加空行
        doc.add_paragraph("")  # 添加空行，保证表格与下一个内容的间隔


# 插入带勾选的表格到Word模板
def insert_table_with_checkboxes(doc, data):

    para = doc.add_paragraph()
    run = para.add_run("损伤序列表")
    run.bold = True  # 加粗
    run.font.size = Pt(20)  # 设置大字体
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
    table = doc.add_table(rows=1, cols=12)  # 创建表格

    # 设置表头
    headers = ['序号', '绳号', '损伤位置（m）', '量值（%）', '类型', '直径', '轻微', '轻度', '中度', '较重', '严重', '超限']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # 插入每一行数据
    for index, row in enumerate(data):
        row_cells = table.add_row().cells
        row_cells[0].text = str(index + 1)  # 序号
        row_cells[1].text = str(row['position'])  # 绳号
        row_cells[2].text = str(row['position'])  # 损伤位置（m）
        row_cells[3].text = f"{row['LMA_value'] * 100:.2f}"  # 量值（%）
        row_cells[4].text = str(row['defect_type'])  # 类型
        row_cells[5].text = str(row['diameter'])  # 直径

        # 根据LMA值打勾
        checks = classify_and_check_lma(row['LMA_value'])
        for i, check in enumerate(checks):
            if check:
                row_cells[6 + i].text = '✔'  # 打勾

        # 设置表格边框为实线，黑色边框
        for row in table.rows:
            for cell in row.cells:
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run()
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)

                tc_pr = cell._element.get_or_add_tcPr()  # 获取单元格的属性
                borders = tc_pr.find(qn('w:tcBorders'))  # 查找边框元素
                if borders is None:
                    borders = OxmlElement('w:tcBorders')  # 如果没有找到边框元素，创建一个新的
                    tc_pr.append(borders)  # 将新的边框元素添加到单元格属性中

                for border in ['top', 'left', 'bottom', 'right']:
                    border_elem = OxmlElement(f'w:{border}')
                    border_elem.set(qn('w:val'), 'single')  # 设置为实线
                    border_elem.set(qn('w:space'), '0')  # 无间隔
                    border_elem.set(qn('w:sz'), '4')  # 设置线宽
                    border_elem.set(qn('w:color'), '000000')  # 黑色边框
                    borders.append(border_elem)

    doc.add_paragraph("")  # 空行


def report_pdf_Rope_Issue(date_to_query):
    # 获取指定日期的数据
    session = Session()
    data = get_event_with_picture_data(session, date_to_query)
    session.close()
    if data:
        # 设置模板路径和输出路径
        template_path = 'templates.docx'  # Word模板路径
        output_path = f'..\FlaskCode\output_with_tables_{date_to_query}.docx'  # 输出文件路径

        # 打开Word模板并插入查询到的数据
        doc = Document(template_path)
        insert_table_with_images(doc, data)  # 插入带图片的表格
        doc.add_paragraph("")  # 空行
        doc.add_paragraph("")  # 空行
        doc.add_paragraph("")  # 空行
        doc.add_paragraph("")  # 空行
        doc.add_paragraph("")  # 空行
        insert_table_with_checkboxes(doc, data)  # 插入带勾选的表格
        doc.save(output_path)
        print(f"文件已保存: {output_path}")
        return output_path
    else:
        print(f"没有找到 {date_to_query} 的数据！")
        return